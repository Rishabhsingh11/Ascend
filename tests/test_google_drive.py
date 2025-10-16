# tests/test_google_drive_service_account.py
"""Test script to verify Google Drive access using Service Account."""

import os
import io
from pathlib import Path
import PyPDF2
import docx

from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload


# Service Account scopes
SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/spreadsheets'
]


def authenticate_google_drive():
    """Authenticate with Google Drive API using Service Account."""
    print("\n🔐 Authenticating with Google Drive (Service Account)...")
    
    creds_path = "credentials/credentials.json"
    
    if not os.path.exists(creds_path):
        raise FileNotFoundError(f"❌ {creds_path} not found!")
    
    try:
        creds = Credentials.from_service_account_file(
            creds_path,
            scopes=SCOPES
        )
        
        service = build('drive', 'v3', credentials=creds)
        
        print(f"✅ Successfully authenticated with Service Account")
        print(f"   Email: {creds.service_account_email}")
        
        return service
        
    except Exception as e:
        print(f"❌ Authentication failed: {e}")
        raise


def find_folder_by_name(service, folder_name):
    """Find a folder ID by folder name."""
    print(f"\n🔍 Searching for folder: '{folder_name}'...")
    
    try:
        query = f"mimeType='application/vnd.google-apps.folder' and name='{folder_name}' and trashed=false"
        
        results = service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name, parents)',
            pageSize=10
        ).execute()
        
        folders = results.get('files', [])
        
        if not folders:
            print(f"❌ Folder '{folder_name}' not found")
            print("\n⚠️  Important: Service accounts can only see files/folders that are explicitly shared with them!")
            print(f"\nTo fix this:")
            print(f"1. Go to your Google Drive")
            print(f"2. Right-click on the '{folder_name}' folder")
            print(f"3. Click 'Share'")
            print(f"4. Add this email as an Editor:")
            
            # Try to get service account email
            try:
                about = service.about().get(fields='user').execute()
                print(f"   {about['user']['emailAddress']}")
            except:
                print(f"   (Check your credentials.json for 'client_email')")
            
            return None
        
        if len(folders) > 1:
            print(f"⚠️  Found {len(folders)} folders with name '{folder_name}':")
            for idx, folder in enumerate(folders, 1):
                print(f"   {idx}. {folder['name']} (ID: {folder['id']})")
            print(f"\n✅ Using the first one: {folders[0]['id']}")
        else:
            print(f"✅ Found folder: {folders[0]['name']} (ID: {folders[0]['id']})")
        
        return folders[0]['id']
        
    except Exception as e:
        print(f"❌ Error finding folder: {e}")
        return None


def list_resume_files_in_folder(service, folder_id, folder_name):
    """List all PDF and DOCX files in a specific folder."""
    print(f"\n📁 Searching for resume files in '{folder_name}' folder...")
    
    try:
        query = f"(mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') and '{folder_id}' in parents and trashed=false"
        
        results = service.files().list(
            q=query,
            spaces='drive',
            fields="files(id, name, mimeType, size, modifiedTime)",
            pageSize=50
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            print(f"❌ No PDF or DOCX files found in '{folder_name}' folder")
            return []
        
        print(f"\n✅ Found {len(files)} resume file(s) in '{folder_name}':\n")
        print("=" * 80)
        
        for idx, file in enumerate(files, 1):
            size_kb = int(file.get('size', 0)) / 1024 if file.get('size') else 0
            file_type = "PDF" if "pdf" in file['mimeType'] else "DOCX"
            print(f"{idx}. {file['name']}")
            print(f"   Type: {file_type} | Size: {size_kb:.1f} KB")
            print(f"   ID: {file['id']}")
            print(f"   Modified: {file.get('modifiedTime', 'N/A')}")
            print("-" * 80)
        
        return files
        
    except Exception as e:
        print(f"❌ Error listing files: {e}")
        return []


def download_file(service, file_id, file_name):
    """Download a file from Google Drive."""
    print(f"\n⬇️  Downloading: {file_name}")
    
    try:
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request, chunksize=204800)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
            if status:
                progress = int(status.progress() * 100)
                print(f"   Progress: {progress}%", end='\r')
        
        print(f"\n✅ Download complete")
        
        # Save to disk
        fh.seek(0)
        with open(file_name, 'wb') as f:
            f.write(fh.read())
        
        print(f"✅ Saved to: {file_name}")
        return True
        
    except Exception as e:
        print(f"❌ Error downloading file: {e}")
        return False


def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    print(f"\n📄 Extracting text from PDF...")
    
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            print(f"   Pages: {num_pages}")
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                text += page_text + "\n"
                print(f"   Extracted page {page_num}/{num_pages}", end='\r')
        
        print(f"\n✅ Extracted {len(text)} characters")
        return text.strip()
        
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return None


def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    print(f"\n📄 Extracting text from DOCX...")
    
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        print(f"✅ Extracted {len(text)} characters")
        return text.strip()
        
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return None


def display_text_preview(text, max_chars=500):
    """Display a preview of the extracted text."""
    print("\n" + "=" * 80)
    print("📝 TEXT CONTENT PREVIEW (First 500 characters)")
    print("=" * 80)
    
    if text:
        preview = text[:max_chars]
        print(preview)
        
        if len(text) > max_chars:
            print(f"\n... ({len(text) - max_chars} more characters)")
        
        print("\n" + "=" * 80)
        print(f"📊 Full text length: {len(text)} characters")
        print(f"📊 Number of lines: {len(text.splitlines())}")
        print(f"📊 Number of words: {len(text.split())}")
    else:
        print("❌ No text content to display")


def main():
    """Main test function."""
    print("\n" + "=" * 80)
    print("🧪 GOOGLE DRIVE SERVICE ACCOUNT TEST - ASCEND_ROOT FOLDER")
    print("=" * 80)
    
    # Configuration
    FOLDER_NAME = "Ascend_Root"
    
    # Step 1: Authenticate
    try:
        service = authenticate_google_drive()
    except Exception as e:
        print(f"\n❌ Authentication failed: {e}")
        print("\nPlease ensure:")
        print("1. credentials/credentials.json is a SERVICE ACCOUNT key (not OAuth)")
        print("2. The file contains 'type': 'service_account'")
        print("3. Google Drive API is enabled")
        return
    
    # Step 2: Find the Ascend_Root folder
    folder_id = find_folder_by_name(service, FOLDER_NAME)
    
    if not folder_id:
        return
    
    # Step 3: List files in the folder
    files = list_resume_files_in_folder(service, folder_id, FOLDER_NAME)
    
    if not files:
        print(f"\n⚠️  No files to test in '{FOLDER_NAME}' folder.")
        print(f"Please upload resume files (PDF or DOCX) to the '{FOLDER_NAME}' folder.")
        return
    
    # Step 4: Select file to test
    print("\n" + "=" * 80)
    if len(files) == 1:
        selected_idx = 0
        print(f"🎯 Auto-selecting the only file: {files[0]['name']}")
    else:
        while True:
            try:
                choice = input(f"\nSelect a file to test (1-{len(files)}): ")
                selected_idx = int(choice) - 1
                if 0 <= selected_idx < len(files):
                    break
                print(f"❌ Please enter a number between 1 and {len(files)}")
            except ValueError:
                print("❌ Please enter a valid number")
            except KeyboardInterrupt:
                print("\n\n⚠️  Test cancelled by user")
                return
    
    selected_file = files[selected_idx]
    
    # Step 5: Download file
    success = download_file(
        service,
        selected_file['id'],
        selected_file['name']
    )
    
    if not success:
        return
    
    # Step 6: Extract text
    file_ext = Path(selected_file['name']).suffix.lower()
    
    if file_ext == '.pdf':
        text = extract_text_from_pdf(selected_file['name'])
    elif file_ext in ['.docx', '.doc']:
        text = extract_text_from_docx(selected_file['name'])
    else:
        print(f"❌ Unsupported file type: {file_ext}")
        return
    
    # Step 7: Display preview
    if text:
        display_text_preview(text)
        
        # Optionally save full text
        save = input("\n💾 Save full text to file? (y/n): ").lower()
        if save == 'y':
            output_file = f"{selected_file['name']}_extracted.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"✅ Full text saved to: {output_file}")
    
    # Cleanup downloaded file
    cleanup = input("\n🗑️  Delete downloaded file? (y/n): ").lower()
    if cleanup == 'y':
        try:
            os.remove(selected_file['name'])
            print(f"✅ Deleted: {selected_file['name']}")
        except Exception as e:
            print(f"❌ Could not delete file: {e}")
    
    print("\n" + "=" * 80)
    print("✅ TEST COMPLETE!")
    print("=" * 80)
    print("\nYour Google Drive Service Account setup is working correctly! ✨")
    print(f"Files from '{FOLDER_NAME}' folder are accessible.")
    print("\nYou can now proceed with building the Google Sheets exporter!")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️  Test interrupted by user")
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        import traceback
        traceback.print_exc()
