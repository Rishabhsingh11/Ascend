"""Resume text extraction from various file formats."""

import PyPDF2
import docx
from typing import Optional
from pathlib import Path


class ResumeTextExtractor:
    """Extract text content from resume files."""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If PDF extraction fails
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                print(f"Extracting text from {num_pages} pages...")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"Processed page {page_num}/{num_pages}")
                    
        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If DOCX extraction fails
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            print(f"Extracted text from {len(doc.paragraphs)} paragraphs")
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")
        
        return text.strip()
    
    def extract_text(self, file_path: str, mime_type: Optional[str] = None) -> str:
        """Extract text based on file type.
        
        Args:
            file_path: Path to the resume file
            mime_type: Optional MIME type hint
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is unsupported
        """
        file_ext = Path(file_path).suffix.lower()
        
        if mime_type == 'application/pdf' or file_ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ] or file_ext in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type or file_ext}")
